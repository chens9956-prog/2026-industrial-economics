import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_run_font(run, ascii_font="Times New Roman", east_asia_font="SimSun"):
    run.font.name = ascii_font
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), east_asia_font)
    rPr.append(rFonts)

def create_hard_return_demo():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. 标题 (硬回车 ↵)
    p_heading = doc.add_paragraph()
    r_h = p_heading.add_run("2. 研究难点及解决策略")
    set_run_font(r_h, ascii_font="Times New Roman", east_asia_font="SimSun")
    r_h.font.size = Pt(14)
    r_h.font.bold = True
    r_h.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # 2. 难点一子标题 (硬回车 ↵)
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.first_line_indent = Pt(22) # 首行缩进
    r_s = p_sub.add_run("难点一：多源异构数据的获取、匹配与“经济-环境协同”指标体系的科学量化。")
    set_run_font(r_s, ascii_font="Times New Roman", east_asia_font="SimSun")
    r_s.font.size = Pt(11)

    # 3. 详细描述 (硬回车 ↵)
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.first_line_indent = Pt(22) # 首行缩进
    p_desc.paragraph_format.line_spacing = 1.25
    r_d = p_desc.add_run("研究难点：“东数西算”涉及算力数据、企业微观碳排放、区域经济运行等多种跨界数据，且“数据流”与“算力流”具有无形性，导致核心变量的代理指标难以直接获取；同时，如何科学赋权并构建兼顾“数字经济活力”与“生态环境质量”的复合协同发展指标是一大挑战。")
    set_run_font(r_d, ascii_font="Times New Roman", east_asia_font="SimSun")
    r_d.font.size = Pt(11)

    output_path = "硬回车段落换行统一示范文档.docx"
    doc.save(output_path)
    print(f"Successfully generated hard return demo doc: {output_path}")

if __name__ == "__main__":
    create_hard_return_demo()
