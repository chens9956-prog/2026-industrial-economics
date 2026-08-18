import os
import json
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import win32com.client

MCKINSEY_BLUE = RGBColor(10, 42, 94) # #0A2A5E
DARK_GRAY = RGBColor(51, 51, 51) # #333333

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "FF0000", "space": "0"},
        bottom={"sz": 12, "color": "FF0000", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def extract_answers():
    transcript_path = r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\logs\transcript.jsonl"
    with open(transcript_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    ans_list = []
    for line in lines:
        try:
            data = json.loads(line)
            if data.get('type') == 'MCP_TOOL':
                content = data.get('content', '')
                ans = ''
                if '{"status":"success"' in content and '"answer"' in content:
                    ans = json.loads(content[content.find('{"status":"success"'):]).get('answer', '')
                elif 'The output was large and was saved to: file://' in content:
                    file_path = content.split('file:///')[1].strip().replace('/', '\\')
                    if os.path.exists(file_path):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            file_content = f.read()
                            if '{"status":"success"' in file_content and '"answer"' in file_content:
                                ans = json.loads(file_content[file_content.find('{"status":"success"'):]).get('answer', '')
                if ans:
                    ans_list.append(ans)
        except:
            pass

    answers = {}
    last_10 = ans_list[-10:]
    for i, a in enumerate(last_10):
        answers[i + 1] = a
    return answers

def create_pdfs():
    output_dir = r"I:\4产业经济学\麦肯锡蓝_要点PDF"
    os.makedirs(output_dir, exist_ok=True)
    
    answers = extract_answers()
    
    if not answers:
        print("未找到任何摘要结果。")
        return
        
    print(f"找到 {len(answers)} 个章节的摘要。")
        
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    for ch_num, answer in answers.items():
        try:
            doc = Document()
            
            # Setup Page layout (A4)
            section = doc.sections[0]
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)
            
            ch_str = f"CH{ch_num:02d}"
            
            # Header Table to hold Title and a Bottom border
            table = doc.add_table(rows=1, cols=1)
            table.autofit = True
            cell = table.cell(0, 0)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.add_run(f"产业经济学 {ch_str} 核心精要")
            run.font.size = Pt(26)
            run.font.name = 'Microsoft YaHei'
            run.font.bold = True
            run.font.color.rgb = MCKINSEY_BLUE
            
            # Add bottom border to the cell to act as an accent line
            set_cell_border(
                cell,
                bottom={"sz": 12, "val": "single", "color": "0A2A5E", "space": "0"}
            )
            
            # Add some space after the header
            doc.add_paragraph()
            
            points = []
            for p_text in answer.split('\n'):
                p_text = p_text.strip()
                if p_text:
                    clean_p = p_text.replace("**", "")
                    points.append(clean_p)
            
            # Output Points
            for point in points:
                # Check if starts with number
                m = re.match(r"^(\d+\.)\s*(.*)", point)
                
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(18)
                para.paragraph_format.line_spacing = 1.3
                
                if m:
                    num_str = m.group(1)
                    content_str = m.group(2)
                    
                    r_num = para.add_run(num_str + " ")
                    r_num.font.size = Pt(12)
                    r_num.font.name = 'Microsoft YaHei'
                    r_num.font.bold = True
                    r_num.font.color.rgb = MCKINSEY_BLUE
                    
                    r_content = para.add_run(content_str)
                    r_content.font.size = Pt(12)
                    r_content.font.name = 'Microsoft YaHei'
                    r_content.font.bold = True
                    r_content.font.color.rgb = DARK_GRAY
                else:
                    r_content = para.add_run(point)
                    r_content.font.size = Pt(12)
                    r_content.font.name = 'Microsoft YaHei'
                    r_content.font.bold = True
                    r_content.font.color.rgb = DARK_GRAY
                    
            docx_filename = f"产业经济学{ch_str}_极简要点_麦肯锡蓝.docx"
            pdf_filename = f"产业经济学{ch_str}_极简要点_麦肯锡蓝.pdf"
            
            docx_path = os.path.join(output_dir, docx_filename)
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            doc.save(docx_path)
            
            abs_docx_path = os.path.abspath(docx_path)
            abs_pdf_path = os.path.abspath(pdf_path)
            
            word_doc = word.Documents.Open(abs_docx_path)
            word_doc.SaveAs(abs_pdf_path, FileFormat=17)
            word_doc.Close()
            
            os.remove(abs_docx_path)
            
            print(f"Successfully generated {pdf_filename}")
            
        except Exception as e:
            print(f"Failed to process CH{ch_num:02d}: {str(e)}")

    word.Quit()
    print("All PDFs processed.")

if __name__ == "__main__":
    create_pdfs()
