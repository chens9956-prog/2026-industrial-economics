import os
import json
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import win32com.client

# Precision Blueprint Colors
BLUEPRINT_BLUE = RGBColor(0, 112, 192) # #0070C0 (Technical blueprint blue)
DARK_TEXT = RGBColor(0, 0, 0) # Pure black for high precision unbolded text

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_run_font(run, font_name, font_size, color, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    # Extremely important for Chinese text in python-docx to avoid fallback font changes
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def extract_answers():
    transcript_path = r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\logs\transcript.jsonl"
    with open(transcript_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    ans_list = []
    for line in lines:
        try:
            data = json.loads(line)
            if data.get('type') == 'MCP_TOOL':
                content = data.get('content', '')
                ans = ''
                if '{"status":"success"' in content and '"answer"' in content:
                    ans = json.loads(content[content.find('{"status":"success"'):]).get('answer', '')
                elif 'The output was large and was saved to: file://' in content:
                    file_path = content.split('file:///')[1].strip().replace('/', '\\')
                    if os.path.exists(file_path):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            file_content = f.read()
                            if '{"status":"success"' in file_content and '"answer"' in file_content:
                                ans = json.loads(file_content[file_content.find('{"status":"success"'):]).get('answer', '')
                if ans:
                    ans_list.append(ans)
        except:
            pass

    answers = {}
    last_10 = ans_list[-10:]
    for i, a in enumerate(last_10):
        answers[i + 1] = a
    return answers

def create_pdfs():
    output_dir = r"I:\4产业经济学\工程蓝图_要点PDF"
    os.makedirs(output_dir, exist_ok=True)
    
    answers = extract_answers()
    
    if not answers:
        print("未找到任何摘要结果。")
        return
        
    print(f"找到 {len(answers)} 个章节的摘要。")
        
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    for ch_num, answer in answers.items():
        try:
            doc = Document()
            
            # Setup Page layout (A4)
            section = doc.sections[0]
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            # Reduce top/bottom margins slightly to ensure it fits on one page
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)
            
            ch_str = f"CH{ch_num:02d}"
            
            # Header Table
            table = doc.add_table(rows=1, cols=1)
            table.autofit = True
            cell = table.cell(0, 0)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.add_run(f"产业经济学 {ch_str} 知识要点汇编")
            set_run_font(run, 'Microsoft YaHei', 24, BLUEPRINT_BLUE, bold=False)
            
            set_cell_border(
                cell,
                bottom={"sz": 18, "val": "double", "color": "0070C0", "space": "0"}
            )
            
            doc.add_paragraph()
            
            points = []
            for p_text in answer.split('\n'):
                p_text = p_text.strip()
                if p_text:
                    # Remove markdown bold
                    clean_p = p_text.replace("**", "")
                    # Remove citation numbers like [1], [6, 7], [1,2,3]
                    clean_p = re.sub(r'\[\s*\d+(?:\s*,\s*\d+)*\s*\]', '', clean_p)
                    # Also handle edge cases where there might be a trailing dot right after citation 
                    # by trimming whitespace before punctuation if necessary, but just stripping citations is enough
                    points.append(clean_p)
            
            for point in points:
                m = re.match(r"^(\d+\.)\s*(.*)", point)
                
                para = doc.add_paragraph()
                # Reduce spacing to guarantee fit on one page
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.line_spacing = 1.2
                
                # Use 11pt font to further guarantee fit
                base_size = 11
                
                if m:
                    num_str = m.group(1)
                    content_str = m.group(2)
                    
                    r_num = para.add_run(num_str + " ")
                    set_run_font(r_num, 'Microsoft YaHei', base_size, BLUEPRINT_BLUE, bold=False)
                    
                    r_content = para.add_run(content_str)
                    set_run_font(r_content, 'Microsoft YaHei', base_size, DARK_TEXT, bold=False)
                else:
                    r_content = para.add_run(point)
                    set_run_font(r_content, 'Microsoft YaHei', base_size, DARK_TEXT, bold=False)
                    
            docx_filename = f"产业经济学{ch_str}_要点汇编_工程蓝图.docx"
            pdf_filename = f"产业经济学{ch_str}_要点汇编_工程蓝图.pdf"
            
            docx_path = os.path.join(output_dir, docx_filename)
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            doc.save(docx_path)
            
            abs_docx_path = os.path.abspath(docx_path)
            abs_pdf_path = os.path.abspath(pdf_path)
            
            word_doc = word.Documents.Open(abs_docx_path)
            word_doc.SaveAs(abs_pdf_path, FileFormat=17)
            word_doc.Close()
            
            os.remove(abs_docx_path)
            
            print(f"Successfully generated {pdf_filename}")
            
        except Exception as e:
            print(f"Failed to process CH{ch_num:02d}: {str(e)}")

    word.Quit()
    print("All PDFs processed.")

if __name__ == "__main__":
    create_pdfs()
