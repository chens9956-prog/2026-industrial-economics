import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_run_font(run, ascii_font="Times New Roman", east_asia_font="SimSun"):
    run.font.name = ascii_font
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), east_asia_font)
    rPr.append(rFonts)

def generate_mindmap_word():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. 标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("《生态文明和生态安全：人与自然共生演化理论》\n全书四层架构思维导图")
    set_run_font(r_title, ascii_font="Times New Roman", east_asia_font="SimSun")
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("原著作者：张智光 教授（南京林业大学）  |  导图整理：Antigravity AI  |  2026年8月")
    set_run_font(r_sub, ascii_font="Times New Roman", east_asia_font="SimSun")
    r_sub.font.size = Pt(10.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 结构说明框
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F2F5F8")
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    rt = p.add_run("📌 四层思维导图结构说明：\n")
    set_run_font(rt, ascii_font="Times New Roman", east_asia_font="SimSun")
    rt.font.bold = True
    rt.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    rc = p.add_run("• 第一层（中心节点）：专著核心主题《生态文明和生态安全：人与自然共生演化理论》；\n• 第二层（一级分支）：三大篇章（第一篇 总论 / 第二篇 生态文明论 / 第三篇 生态安全论）；\n• 第三层（二级分支）：全书共 21 个核心章节系统展开；\n• 第四层（三级分支）：各章节最核心的方法论、机理模型与决策建议。")
    set_run_font(rc, ascii_font="Times New Roman", east_asia_font="SimSun")
    rc.font.size = Pt(9.5)
    rc.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph() # 空行

    # 思维导图结构数据
    parts_data = [
        ("第一篇 总论：理论基础与总体框架构筑", [
            ("第一章 生态文明与生态安全的理论基础", [
                "概念界定（生态文明内涵、生态安全构成、两者“共生-安全”耦合关联）",
                "哲学与经济学基础（浅层/深层生态学、文明形态史观、循环经济与低碳经济）",
                "生态学与生态经济学基础（共生理论、生态系统健康、资源-环境-经济复合系统）",
                "管理学与社会学基础（生态系统服务价值评估、环境风险评价、生态文化理论）"
            ]),
            ("第二章 我国生态文明与生态安全的现状与问题", [
                "自然环境子系统诊断（土地、水域、大气、生物多样性与外来物种入侵）",
                "人类活动子系统诊断（生态经济、生态科技、生态法律制度与行为文明）",
                "林业生态调控影响（森林对水源涵养、大气污染物沉降与重金属修复作用）"
            ]),
            ("第三章 总体理论构筑：人与自然共生演化理论", [
                "“人-自然-经济-社会”四维共生演化理论框架起源与逻辑起点",
                "复合共生系统平衡与演进模型（Symbiotic Evolution Model）",
                "微观-中观-宏观三级共生运行机制与协同路径"
            ]),
            ("第四章 微观-企业层共生：企业绿色科技创新", [
                "企业绿色创新的“意愿-行为-绩效”（WBP）三阶驱动模型",
                "环境规制与市场激励双重驱动下的企业绿色响应机理",
                "企业生态效益与经济效益“双赢”（Win-Win）实现机制"
            ]),
            ("第五章 中观-供应链共生：绿色共生型供应链", [
                "绿色共生型供应链内在驱动力与上下游节点协同博弈",
                "物质循环与能量梯级利用的产业链生态化重构",
                "共生供应链绩效评价指标体系与生态风险管控"
            ]),
            ("第六章 宏观-全社会共生：超循环经济理论与实例", [
                "超循环经济（Hypercyclic Economy）自组织与正反馈理论模型",
                "社会-经济-生态三大子系统的大循环耦合机制",
                "国内外超循环经济示范区典型案例与经验借鉴"
            ])
        ]),
        ("第二篇 生态文明论：概念界定、测度模型与时空演化", [
            ("第七章 生态文明国内外研究综述", [
                "从环境保护到“文明形态”研究演进脉络综述",
                "国内外生态文明评价指标体系对比分析",
                "现有测度方法的创新突破点与未尽议题"
            ]),
            ("第八章 生态文明观：新时代的发展观", [
                "生态文明观的核心价值观与导向演进",
                "“绿水青山就是金山银山”两山转化理论逻辑",
                "生态文明在国家现代化建设中的战略定位"
            ]),
            ("第九章 生态文明“阈值”与“绿值”二步测度方法", [
                "“二步测度法”（Two-Step Measurement Method）创新思路与数学模型",
                "临界“阈值”（Threshold）判定与生态安全下限标准",
                "发展“绿值”（Green Value）综合量化与动态权重构建"
            ]),
            ("第十章 生态文明测度的 PSIR-SEM 建模", [
                "PSIR（压力-状态-影响-响应）扩展框架构建",
                "结构方程模型（SEM）在潜变量量化中的应用",
                "与传统熵权法/主成分分析法的优劣比较与检验"
            ]),
            ("第十一章 中国生态文明测度及其时空演化规律分析", [
                "全国及 31 省域生态文明二步测度实证结果",
                "时空演化特征（东中西部梯度演进与区域收敛趋势）",
                "生态文明发展阶段分类判定与政策契合度"
            ]),
            ("第十二章 林业对生态文明的贡献及其“结构微笑曲线”", [
                "林业在生态文明建设中的基础与枢纽地位",
                "林业产业链“结构微笑曲线”（Structure Smile Curve）模型",
                "前端生态保育与后端生态旅游高附加值段提升路径"
            ]),
            ("第十三章 中国生态文明状况的根源回溯与对策建议", [
                "制度缺失、市场失灵与行为滞后的深层根源剖析",
                "制度创新、绿色科技与全民参与三位一体对策",
                "面向 2035/2050 年生态文明建设路线图与政策组合拳"
            ])
        ]),
        ("第三篇 生态安全论：安全指数、空间测度与屏障构筑", [
            ("第十四章 生态安全的国内外研究综述", [
                "生态安全概念演变、内涵界定与预警机制综述",
                "区域与产业生态安全评价方法对比",
                "生态安全研究前沿趋势与共生范式转向"
            ]),
            ("第十五章 生态安全指数的共生空间测度方法", [
                "共生空间（Symbiotic Space）概念维度与几何表征",
                "生态安全指数共生空间测度模型与算法设计",
                "空间溢出效应与跨区域协同安全机制"
            ]),
            ("第十六章 生态安全测度的 PSIR-SEM 建模", [
                "生态安全 PSIR-SEM 结构方程潜变量设计",
                "模型路径系数拟合与安全因果链条验证",
                "关键敏感因子识别与早期预警阈值触发机制"
            ]),
            ("第十七章 中国生态安全指数测度及其时空演化（以林业为例）", [
                "国家与省域林业生态安全指数时间序列测算",
                "生态安全屏障空间格局演进与区域分化轨迹",
                "林业生态安全预警等级划分与动态调控"
            ]),
            ("第十八章 林业产业的生态安全性评估", [
                "木材加工与林产化工的资源环境负荷评估",
                "林业产业生态安全性综合评价指标体系",
                "传统林产工业的清洁生产与绿色转型路径"
            ]),
            ("第十九章 森林旅游业的生态安全性评估", [
                "森林生态旅游的环境承载力（Carrying Capacity）测算",
                "旅游开发强度与生态保护的演化博弈模型",
                "低碳森林旅游模式与生态预警控制机制"
            ]),
            ("第二十章 森林食品业的生态安全性评估", [
                "森林食品产地环境质量与土壤/水质安全评估",
                "“从森林到餐桌”全产业链生态安全保障体系",
                "森林有机食品品牌价值与生态产品价值实现"
            ]),
            ("第二十一章 中国生态安全屏障的构筑", [
                "“三区四带”国家生态安全屏障总体战略布局",
                "跨区域生态补偿与安全共建联防联控机制",
                "面向未来的国家生态安全长效保障战略与制度安排"
            ])
        ])
    ]

    for p_title_text, chapters in parts_data:
        h1 = doc.add_heading(level=1)
        r_h1 = h1.add_run(p_title_text)
        set_run_font(r_h1, ascii_font="Times New Roman", east_asia_font="SimSun")
        r_h1.font.size = Pt(14)
        r_h1.font.bold = True
        r_h1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        for c_title_text, points in chapters:
            h2 = doc.add_heading(level=2)
            r_h2 = h2.add_run(c_title_text)
            set_run_font(r_h2, ascii_font="Times New Roman", east_asia_font="SimSun")
            r_h2.font.size = Pt(12)
            r_h2.font.bold = True
            r_h2.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)

            for pt in points:
                p_item = doc.add_paragraph(style='List Bullet')
                p_item.paragraph_format.space_before = Pt(2)
                p_item.paragraph_format.space_after = Pt(3)
                p_item.paragraph_format.line_spacing = 1.25
                r_pt = p_item.add_run(pt)
                set_run_font(r_pt, ascii_font="Times New Roman", east_asia_font="SimSun")
                r_pt.font.size = Pt(10)

        doc.add_paragraph() # 空行

    output_path = "张智光_生态文明和生态安全_4层思维导图.docx"
    doc.save(output_path)
    print(f"Successfully generated Mindmap Word doc: {output_path}")

if __name__ == "__main__":
    generate_mindmap_word()
