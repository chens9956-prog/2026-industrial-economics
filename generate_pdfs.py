import os
import json
from docx import Document
from docx.shared import Pt
import win32com.client

def extract_answers():
    transcript_path = r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\logs\transcript.jsonl"
    with open(transcript_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    ans_list = []
    for line in lines:
        try:
            data = json.loads(line)
            if data.get('type') == 'MCP_TOOL':
                content = data.get('content', '')
                ans = ''
                if '{"status":"success"' in content and '"answer"' in content:
                    ans = json.loads(content[content.find('{"status":"success"'):]).get('answer', '')
                elif 'The output was large and was saved to: file://' in content:
                    file_path = content.split('file:///')[1].strip().replace('/', '\\')
                    if os.path.exists(file_path):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            file_content = f.read()
                            if '{"status":"success"' in file_content and '"answer"' in file_content:
                                ans = json.loads(file_content[file_content.find('{"status":"success"'):]).get('answer', '')
                if ans:
                    ans_list.append(ans)
        except:
            pass

    answers = {}
    last_10 = ans_list[-10:]
    for i, a in enumerate(last_10):
        answers[i + 1] = a
    return answers

def create_pdfs():
    output_dir = r"I:\4产业经济学\章节摘要"
    os.makedirs(output_dir, exist_ok=True)
    
    answers = extract_answers()
    
    if not answers:
        print("未找到任何摘要结果。")
        return
        
    print(f"找到 {len(answers)} 个章节的摘要。")
        
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    for ch_num, answer in answers.items():
        try:
            doc = Document()
            ch_str = f"CH{ch_num:02d}"
            doc.add_heading(f"产业经济学 {ch_str} 核心摘要", 0)
            
            for p in answer.split('\n'):
                p = p.strip()
                if not p:
                    continue
                
                # Remove any existing markdown bold asterisks to clean the text
                clean_p = p.replace("**", "")
                
                para = doc.add_paragraph()
                run = para.add_run(clean_p)
                run.font.size = Pt(12)
                run.font.name = 'Microsoft YaHei'
                # Bold the entire line for high visibility
                run.bold = True
            
            docx_filename = f"产业经济学{ch_str}_精炼摘要.docx"
            pdf_filename = f"产业经济学{ch_str}_精炼摘要.pdf"
            
            docx_path = os.path.join(output_dir, docx_filename)
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            doc.save(docx_path)
            
            abs_docx_path = os.path.abspath(docx_path)
            abs_pdf_path = os.path.abspath(pdf_path)
            
            word_doc = word.Documents.Open(abs_docx_path)
            word_doc.SaveAs(abs_pdf_path, FileFormat=17)
            word_doc.Close()
            
            os.remove(abs_docx_path)
            
            print(f"Successfully generated {pdf_filename}")
            
        except Exception as e:
            print(f"Failed to process CH{ch_num:02d}: {str(e)}")

    word.Quit()
    print("All PDFs processed.")

if __name__ == "__main__":
    create_pdfs()
