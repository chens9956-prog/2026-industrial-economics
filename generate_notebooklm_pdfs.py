import os
import json
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import win32com.client

# NotebookLM Colors
NBLM_TITLE = RGBColor(32, 33, 36) # Google Dark Gray
NBLM_TEXT = RGBColor(60, 64, 67) # Google Subtext Gray
NBLM_ACCENT = "DADCE0" # Google border gray for the quote line

def set_para_border(paragraph, **kwargs):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.first_child_found_in("w:pBdr")
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    for edge in ('top', 'left', 'bottom', 'right', 'between'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = pBdr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                pBdr.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_run_font(run, font_name, font_size, color, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def extract_answers():
    transcript_path = r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\logs\transcript.jsonl"
    with open(transcript_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    ans_list = []
    for line in lines:
        try:
            data = json.loads(line)
            if data.get('type') == 'MCP_TOOL':
                content = data.get('content', '')
                ans = ''
                if '{"status":"success"' in content and '"answer"' in content:
                    ans = json.loads(content[content.find('{"status":"success"'):]).get('answer', '')
                elif 'The output was large and was saved to: file://' in content:
                    file_path = content.split('file:///')[1].strip().replace('/', '\\')
                    if os.path.exists(file_path):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            file_content = f.read()
                            if '{"status":"success"' in file_content and '"answer"' in file_content:
                                ans = json.loads(file_content[file_content.find('{"status":"success"'):]).get('answer', '')
                if ans:
                    ans_list.append(ans)
        except:
            pass

    answers = {}
    last_10 = ans_list[-10:]
    for i, a in enumerate(last_10):
        answers[i + 1] = a
    return answers

def create_pdfs():
    output_dir = r"I:\4产业经济学\NotebookLM风_要点PDF"
    os.makedirs(output_dir, exist_ok=True)
    
    answers = extract_answers()
    if not answers:
        print("未找到任何摘要结果。")
        return
        
    print(f"找到 {len(answers)} 个章节的摘要。")
        
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    for ch_num, answer in answers.items():
        try:
            doc = Document()
            
            section = doc.sections[0]
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)
            
            ch_str = f"CH{ch_num:02d}"
            
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_title.add_run(f"✨ 产业经济学 {ch_str} 核心精要")
            set_run_font(run, 'Microsoft YaHei', 20, NBLM_TITLE, bold=False)
            
            doc.add_paragraph()
            
            points = []
            for p_text in answer.split('\n'):
                p_text = p_text.strip()
                if p_text:
                    clean_p = p_text.replace("**", "")
                    clean_p = re.sub(r'\[\s*\d+(?:\s*,\s*\d+)*\s*\]', '', clean_p)
                    points.append(clean_p)
            
            for point in points:
                m = re.match(r"^(\d+\.)\s*(.*)", point)
                
                para = doc.add_paragraph()
                
                # Quote-centric layout styling
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.line_spacing = 1.2
                
                # Indentation for the quote block effect
                para.paragraph_format.left_indent = Pt(14)
                
                # Left thick border mimicking NotebookLM quote block
                # sz: border thickness (in 1/8 pt), 32 = 4pt
                # space: spacing from text
                set_para_border(
                    para,
                    left={"sz": 32, "val": "single", "color": NBLM_ACCENT, "space": "10"}
                )
                
                base_size = 11
                
                if m:
                    num_str = m.group(1)
                    content_str = m.group(2)
                    
                    r_num = para.add_run(num_str + " ")
                    set_run_font(r_num, 'Microsoft YaHei', base_size, NBLM_TITLE, bold=False)
                    
                    r_content = para.add_run(content_str)
                    set_run_font(r_content, 'Microsoft YaHei', base_size, NBLM_TEXT, bold=False)
                else:
                    r_content = para.add_run(point)
                    set_run_font(r_content, 'Microsoft YaHei', base_size, NBLM_TEXT, bold=False)
                    
            docx_filename = f"产业经济学{ch_str}_引用风_NotebookLM.docx"
            pdf_filename = f"产业经济学{ch_str}_引用风_NotebookLM.pdf"
            
            docx_path = os.path.join(output_dir, docx_filename)
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            doc.save(docx_path)
            
            abs_docx_path = os.path.abspath(docx_path)
            abs_pdf_path = os.path.abspath(pdf_path)
            
            word_doc = word.Documents.Open(abs_docx_path)
            word_doc.SaveAs(abs_pdf_path, FileFormat=17)
            word_doc.Close()
            
            os.remove(abs_docx_path)
            
            print(f"Successfully generated {pdf_filename}")
            
        except Exception as e:
            print(f"Failed to process CH{ch_num:02d}: {str(e)}")

    word.Quit()
    print("All PDFs processed.")

if __name__ == "__main__":
    create_pdfs()
