import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_run_font(run, ascii_font="Times New Roman", east_asia_font="SimSun"):
    run.font.name = ascii_font
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), east_asia_font)
    rPr.append(rFonts)

def create_full_spring_doc():
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 默认样式
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)

    # 1. 大标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("春")
    set_run_font(r_title, ascii_font="Times New Roman", east_asia_font="SimSun")
    r_title.font.size = Pt(24)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(16)
    r_author = p_author.add_run("朱自清 (Zhu Ziqing, 1898–1948)")
    set_run_font(r_author, ascii_font="Times New Roman", east_asia_font="SimSun")
    r_author.font.size = Pt(11)
    r_author.font.italic = True
    r_author.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 规则展示卡片
    rule_table = doc.add_table(rows=1, cols=1)
    rule_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = rule_table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F2F5F8")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    rt = p.add_run("📌 全局字体与标点新规应用说明 (2026年最新规范测试)：\n")
    set_run_font(rt, ascii_font="Times New Roman", east_asia_font="SimSun")
    rt.font.bold = True
    rt.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    rc = p.add_run("• 中文字体：统一使用 宋体 (SimSun)；英文与数字：统一使用 Times New Roman（如 2026 年、Classic Article）。\n• 正文标点：统一使用全角中文标点与全角弯双引号“与”；\n• 特别例外：文末参考文献 (References) 的标点符号统一保留学术格式的英文半角标点。")
    set_run_font(rc, ascii_font="Times New Roman", east_asia_font="SimSun")
    rc.font.size = Pt(9.5)
    rc.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph() # 间距

    # 文章分段内容
    sections_content = [
        ("一、盼春", [
            "盼望着，盼望着，东风来了，春天的脚步近了（2026年春季经典篇章）。",
            "一切都像刚睡醒的样子，欣欣然张开了眼。山朗润起来了，水涨起来了，太阳的脸红起来了。"
        ]),
        ("二、绘春", [
            "小草偷偷地从土里钻出来，嫩嫩的，绿绿的。园子里，田野里，瞧去，一大片一大片满是的。坐着，躺着，打两个滚，踢几脚球，赛几趟跑，捉几迷藏。风轻悄悄的，草软绵绵的。",
            "桃树、杏树、梨树，你不让我，我不让你，都开满了花“赶趟儿”。红的像火，粉的像霞，白的像雪。花里带着甜味儿；闭了眼，树上仿佛已经满是桃儿、杏儿、梨儿。花下成千成百的蜜蜂嗡嗡地闹着，大小的蝴蝶飞来飞去。野花散在草丛里，像眼睛，像星星，还眨呀眨的。",
            "“吹面不寒杨柳风”，不错的，像母亲的手抚摸着你。风里带来些新翻的泥土的气息，混着青草味儿，还有各种花的香，都在微微润湿的空气里孵化。鸟儿将巢安在繁花嫩叶当中，高兴起来了，呼朋引伴地卖弄清脆的喉咙，唱出宛转的曲子，跟轻风流水应和着。牛背上牧童的短笛，这时候也成天嘹亮地响着。",
            "雨是最寻常的，一下就是三两天。可别恼。看，像牛毛，像花针，像细丝，密密地斜织着，人家屋顶上全笼着一层薄烟。树叶儿却绿得发亮，小草儿也青得逼你的眼。傍晚时候，上灯了，一点点黄晕的光，烘托出一片安静而和平的夜。在乡下，小路上，石桥边，有撑起伞慢慢走着的人，地里还有工作的农民，披着蓑戴着笠。他们的房屋，稀稀疏疏的，在雨里静默着。",
            "天上风筝渐渐多了，地上孩子也渐渐多了。城里乡下，家家户户，老老小小，也赶趟儿似的，一个个都出来了。舒活舒活筋骨，抖擞抖擞精神，各做各的一份事去。“一年之计在于春”，刚起头儿，有的是功夫，有的是希望。"
        ]),
        ("三、赞春", [
            "春天像刚落地的娃娃，从头到脚都是新的，它生长着。",
            "春天像小姑娘，花枝招展的，笑着，走着。",
            "春天像健壮的青年，有铁一般的胳膊和腰脚，领着我们向前去。"
        ])
    ]

    for title, paras in sections_content:
        h = doc.add_heading(level=2)
        hr = h.add_run(title)
        set_run_font(hr, ascii_font="Times New Roman", east_asia_font="SimSun")
        hr.font.size = Pt(13)
        hr.font.bold = True
        hr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        for p_text in paras:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25
            pr = p.add_run(p_text)
            set_run_font(pr, ascii_font="Times New Roman", east_asia_font="SimSun")

    doc.add_paragraph() # 间距

    # 文末参考文献 (References) - 标点例外测试
    h_ref = doc.add_heading(level=2)
    hr_ref = h_ref.add_run("参考文献 (References)")
    set_run_font(hr_ref, ascii_font="Times New Roman", east_asia_font="SimSun")
    hr_ref.font.size = Pt(13)
    hr_ref.font.bold = True
    hr_ref.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    references_list = [
        "[1] 朱自清. 朱自清散文全集 [M]. 北京: 人民文学出版社, 1998: 45-48.",
        "[2] Smith, J. A., & Brown, L. M. (2024). Aesthetic Elements in Modern Chinese Prose: A Study of Zhu Ziqing's Works. Journal of Asian Literature, 45(2), 89-104. https://doi.org/10.1016/j.jal.2024.02.001",
        "[3] 高建刚. 产业经济学与现代文学跨学科研究 [M]. 北京: 高等教育出版社, 2026: 120-135."
    ]

    for ref in references_list:
        rp = doc.add_paragraph()
        rp.paragraph_format.space_before = Pt(2)
        rp.paragraph_format.space_after = Pt(4)
        rr = rp.add_run(ref)
        set_run_font(rr, ascii_font="Times New Roman", east_asia_font="SimSun")
        rr.font.size = Pt(9.5)
        rr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    output_path = "朱自清_春_新规效果完整示范文档.docx"
    doc.save(output_path)
    print(f"Successfully generated full demo doc: {output_path}")

if __name__ == "__main__":
    create_full_spring_doc()
