import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_executive_brief():
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 默认字体
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft YaHei'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ================= 1. 报告 Header =================
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = header_table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "1F4E79") # 商务深蓝背景
    set_cell_margins(cell, top=180, bottom=180, left=200, right=200)

    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_sub = p_title.add_run("EXECUTIVE BRIEF | 总裁决策简报\n")
    r_sub.font.name = 'Microsoft YaHei'
    r_sub.font.size = Pt(10)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(0xD9, 0xE1, 0xF2)

    r_main = p_title.add_run("不懂 AI 也能用！两大核心口诀与 AI Agent 超狂自动化指南")
    r_main.font.name = 'Microsoft YaHei'
    r_main.font.size = Pt(16)
    r_main.font.bold = True
    r_main.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 报告元数据表
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(8)
    meta_p.paragraph_format.space_after = Pt(14)
    r_meta = meta_p.add_run("📅 报告日期：2026年8月  |  👤 报告整理：Antigravity AI  |  🎯 报告形式：总裁报告 (Executive Brief)")
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r_meta.font.italic = True

    # ================= Ⅰ. 结论放上面 (Executive Summary) =================
    h1_1 = doc.add_heading(level=1)
    h1_1_run = h1_1.add_run("Ⅰ. 结论与核心决策摘要 (Executive Summary)")
    h1_1_run.font.name = 'Microsoft YaHei'
    h1_1_run.font.size = Pt(14)
    h1_1_run.font.bold = True
    h1_1_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    summary_box = doc.add_table(rows=1, cols=1)
    summary_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_cell = summary_box.cell(0, 0)
    s_cell.width = Inches(6.5)
    set_cell_background(s_cell, "F2F5F8") # 浅蓝灰
    set_cell_margins(s_cell, top=140, bottom=140, left=180, right=180)

    sp = s_cell.paragraphs[0]
    sp.paragraph_format.space_after = Pt(4)
    
    bullets = [
        ("📌 范式转移", "AI 已经从单纯的“对话问答框”演进为能够自主协同、跨应用执行任务的“AI Agent（智能体）”。无需任何 IT 编程背景，管理者与员工仅需掌握两大核心口诀，即可实现工作流 10 倍速自动化。"),
        ("💡 核心口诀一【想到就丢给它】", "打破素材格式限制。随时将会议逐字稿、设计截图、数据表格、公司标准模板丢给 AI，AI 会自动按上下文推演并输出符合预期的高质感成果。"),
        ("💡 核心口诀二【想到就开外挂】", "打破软件壁垒。随时在输入框打上 `@` 挂载 Gmail、Google Drive、Browser 浏览器或 Word 文档生成器，实现多工具实时联动与自动化排程。"),
        ("🚀 管理层启示与价值", "实现“人在通勤，AI 自动干活；人在休息，AI 后台排程”的高效人机协同模式，极大降低企业的行政与沟通摩擦成本，将员工精力集中于高价值决策。")
    ]

    for title, desc in bullets:
        p = s_cell.add_paragraph() if sp != s_cell.paragraphs[0] else sp
        sp = None
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(4)
        rt = p.add_run(f"{title}：")
        rt.font.bold = True
        rt.font.size = Pt(10)
        rt.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        
        rd = p.add_run(desc)
        rd.font.size = Pt(10)
        rd.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph() # 间距

    # ================= Ⅱ. 重点内容放中间 (Key Insights & Detailed Content) =================
    h1_2 = doc.add_heading(level=1)
    h1_2_run = h1_2.add_run("Ⅱ. 重点内容与实践案例拆解 (Key Insights & Core Applications)")
    h1_2_run.font.name = 'Microsoft YaHei'
    h1_2_run.font.size = Pt(14)
    h1_2_run.font.bold = True
    h1_2_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # 重点一
    h2_1 = doc.add_heading(level=2)
    h2_1_run = h2_1.add_run("1. 口诀一【想到就丢给它】：素材与上下文的无缝融合")
    h2_1_run.font.name = 'Microsoft YaHei'
    h2_1_run.font.size = Pt(12)
    h2_1_run.font.bold = True
    h2_1_run.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)

    p1 = doc.add_paragraph()
    p1.add_run("传统 AI 使用的痛点在于用户需要花费大量时间撰写繁琐的 Prompt。而在 Agent 模式下，关键在于直接把已有的上下文素材“丢”给 AI：\n").font.size = Pt(10.5)
    
    p1_a = doc.add_paragraph(style='List Bullet')
    r1_a = p1_a.add_run("案例 A：会议记录与逐字稿自动整理\n")
    r1_a.font.bold = True
    p1_a.add_run("将录音转成的逐字稿文件 + 公司固定的会议记录模板（包含主题、日期、目的、3点重点摘要）同时丢给 AI，AI 即可自动提取核心信息并套用标准格式。")

    p1_b = doc.add_paragraph(style='List Bullet')
    r1_b = p1_b.add_run("案例 B：视觉报表与参考样式复刻\n")
    r1_b.font.bold = True
    p1_b.add_run("在网络上看到设计精美的分析报告时，直接将报告截图丢给 AI 作为参考样式；同时将 Google Drive 里的原始数据丢给 AI，AI 能自动复刻同款排版风格并输出专业图表。")

    # 重点二
    h2_2 = doc.add_heading(level=2)
    h2_2_run = h2_2.add_run("2. 口诀二【想到就开外挂】：生态连动与自动化工作流")
    h2_2_run.font.name = 'Microsoft YaHei'
    h2_2_run.font.size = Pt(12)
    h2_2_run.font.bold = True
    h2_2_run.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)

    p2 = doc.add_paragraph()
    p2.add_run("通过在输入框输入 `@` 或点击 `+` 开启对应应用插件，让 AI Agent 获得直接读取与操作外部工具的能力：\n").font.size = Pt(10.5)

    p2_a = doc.add_paragraph(style='List Bullet')
    r2_a = p2_a.add_run("外挂 A：@Gmail 邮箱自动抽取\n")
    r2_a.font.bold = True
    p2_a.add_run("无需手动下载邮箱附件，直接指令 AI：“寻找 Gmail 中最新的财务会议记录并依照公司模板整理”，AI 自动连接邮箱抓取并完成整理。")

    p2_b = doc.add_paragraph(style='List Bullet')
    r2_b = p2_b.add_run("外挂 B：@Browser + 文档生成器（总裁报告自动产出）\n")
    r2_b.font.bold = True
    p2_b.add_run("开启浏览器插件读取目标商业文章，配合小型文件生成器，AI 自动提炼核心推演逻辑，并直接导出标准“总裁报告”（结论置顶、重点置中、备注置底）格式的 Word (.docx) 文件。")

    p2_c = doc.add_paragraph(style='List Bullet')
    r2_c = p2_c.add_run("外挂 C：自定义 Agent Skills 与排程（Lead Research 商机挖掘）\n")
    r2_c.font.bold = True
    p2_c.add_run("自定义 `.md` 技能规范（设定公司名称、网址、员工数、地区、产业、需求信号与 90 天去重规则），结合【定时排程（Schedule）】功能，实现每周自动挖掘潜在客户或招聘信号。")

    # 重点三：进阶控制机制表格
    h2_3 = doc.add_heading(level=2)
    h2_3_run = h2_3.add_run("3. 提升 Agent 产出质量的三大进阶机制")
    h2_3_run.font.name = 'Microsoft YaHei'
    h2_3_run.font.size = Pt(12)
    h2_3_run.font.bold = True
    h2_3_run.font.color.rgb = RGBColor(0x2F, 0x55, 0x97)

    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["进阶控制机制", "核心功能与作用", "推荐使用策略"]
    for col_idx, text in enumerate(headers):
        c = table.cell(0, col_idx)
        set_cell_background(c, "1F4E79")
        set_cell_margins(c, top=100, bottom=100, left=120, right=120)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    rows_data = [
        ("模型阶层切换\n(Model Tiering)", "依据任务复杂度在右下角切换不同模型能力与 Token 消耗", "高难度跨数据分析选高阶大模型；日常过滤选轻量快速模型"),
        ("侧边栏面板\n(Control Panel)", "点击右上角开启侧边栏，实时监控 AI 读取的文件、模板与视窗", "保持交互透明度，随时追加参考文件或实时纠偏"),
        ("移动端远程控制\n(Remote Control)", "手机端语音下达任务，远程指挥桌面版 AI 运行并产出文件", "利用通勤时间零碎下达指令，上班开机即可直接使用成果")
    ]

    col_widths = [Inches(1.8), Inches(2.5), Inches(2.2)]

    for row_idx, data in enumerate(rows_data, 1):
        for col_idx, text in enumerate(data):
            c = table.cell(row_idx, col_idx)
            c.width = col_widths[col_idx]
            bg_color = "F9FAFB" if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background(c, bg_color)
            set_cell_margins(c, top=100, bottom=100, left=120, right=120)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph() # 间距

    # ================= Ⅲ. 备注放最下 (Notes & Appendix) =================
    h1_3 = doc.add_heading(level=1)
    h1_3_run = h1_3.add_run("Ⅲ. 备注、延伸资源与后续落地行动 (Notes & Next Steps)")
    h1_3_run.font.name = 'Microsoft YaHei'
    h1_3_run.font.size = Pt(14)
    h1_3_run.font.bold = True
    h1_3_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    notes_box = doc.add_table(rows=1, cols=1)
    notes_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    n_cell = notes_box.cell(0, 0)
    n_cell.width = Inches(6.5)
    set_cell_background(n_cell, "FAFAFA") # 浅灰
    set_cell_margins(n_cell, top=140, bottom=140, left=180, right=180)

    np = n_cell.paragraphs[0]
    
    notes_content = [
        ("🔒 1. 数据安全与权限合规备注", "使用 Google Drive、Gmail 或企业内部系统连接外挂时，需确认开启的读取/写入授权范围符合公司信息安全与数据隐私规范。"),
        ("⚡ 2. 自动化排程与去重机制", "在配置如 Lead Research 等定时后台任务时，务必在技能 Prompt 中加入时间窗口过滤规则（如 90 天去重），避免相同商机或数据重复堆积。"),
        ("🔗 3. 来源信息与参考资源", "• 影片标题：《不懂 AI 也能用！兩個口訣享受 AI Agent 超狂自動化！》\n• 讲者出处：汉克蔡 / AI 集（升格创业学院）\n• 影片链接：https://www.youtube.com/watch?v=6AprcceiKIg"),
        ("🎯 4. 管理层落地行动计划 (Next Steps)", "① 梳理部门高频行政/会议流程，建立标准 `.md` 格式的团队专属 Agent Skills 技能库；\n② 统一制定企业内部“总裁报告”与“项目简报”的 Word/PPT 标准模板；\n③ 配置核心业务的自动定时排程，打造无人值守的智能化办公工作流。")
    ]

    for title, desc in notes_content:
        p = n_cell.add_paragraph() if np != n_cell.paragraphs[0] else np
        np = None
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        rt = p.add_run(f"{title}\n")
        rt.font.bold = True
        rt.font.size = Pt(10)
        rt.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        rd = p.add_run(desc)
        rd.font.size = Pt(9.5)
        rd.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    output_path = "总裁报告_AI_Agent超狂自动化与两大口诀解析.docx"
    doc.save(output_path)
    print(f"Successfully generated executive brief Word document at {output_path}")

if __name__ == "__main__":
    create_executive_brief()
