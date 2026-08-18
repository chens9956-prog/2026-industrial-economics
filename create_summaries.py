import os
import json
from docx import Document
from docx.shared import RGBColor, Pt

def create_word_docs():
    output_dir = r"I:\4产业经济学\章节摘要"
    os.makedirs(output_dir, exist_ok=True)
    
    chapter_files = {
        1: r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\steps\949\output.txt",
        2: r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\steps\956\output.txt",
        3: r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\steps\957\output.txt",
        4: r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\steps\960\output.txt",
        5: r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\steps\961\output.txt",
        6: r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\steps\964\output.txt",
        7: r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\steps\965\output.txt",
        8: r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\steps\968\output.txt",
        9: r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\steps\969\output.txt",
        10: r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\steps\970\output.txt"
    }

    for ch_num, json_path in chapter_files.items():
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            data = json.loads(content)
            answer = data.get("answer", "")
            
            doc = Document()
            ch_str = f"CH{ch_num:02d}"
            doc.add_heading(f"产业经济学 {ch_str} 核心摘要", 0)
            
            for p in answer.split('\n'):
                p = p.strip()
                if not p:
                    continue
                
                para = doc.add_paragraph()
                
                # Split by ** to find bold sections
                parts = p.split('**')
                for i, part in enumerate(parts):
                    if not part:
                        continue
                    run = para.add_run(part)
                    if i % 2 == 1:
                        # Inside ** **, make it bold and colored (Navy Blue)
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 51, 153)
                        run.font.size = Pt(12)
                    else:
                        run.font.size = Pt(12)
            
            output_filename = f"产业经济学{ch_str}_精炼摘要.docx"
            output_path = os.path.join(output_dir, output_filename)
            doc.save(output_path)
            
            print(f"Successfully generated {output_filename}")
            
        except Exception as e:
            print(f"Failed to process CH{ch_num:02d}: {str(e)}")

if __name__ == "__main__":
    create_word_docs()
