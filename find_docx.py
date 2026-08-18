import os
import sys
import datetime
from collections import defaultdict

try:
    import docx
except ImportError:
    docx = None

sys.stdout.reconfigure(encoding='utf-8')

KEYWORD = "产业经济学"
DRIVES = ['C:\\', 'D:\\', 'E:\\', 'F:\\', 'G:\\', 'H:\\', 'I:\\', 'J:\\', 'K:\\', 'L:\\']
EXCLUDE_DIRS = {'$RECYCLE.BIN', 'Windows', 'Program Files', 'Program Files (x86)', 'node_modules', '.git', '_Drive_Clean_Backup_'}

def format_size(size_bytes):
    if size_bytes >= 1024 * 1024:
        return f"{size_bytes / (1024 * 1024):.2f} MB"
    elif size_bytes >= 1024:
        return f"{size_bytes / 1024:.2f} KB"
    else:
        return f"{size_bytes} Bytes"

def check_content_contains(filepath):
    if not docx or not filepath.endswith('.docx'):
        return False
    try:
        if os.path.getsize(filepath) > 10 * 1024 * 1024: # >10MB 跳过正文解析
            return False
        doc = docx.Document(filepath)
        for p in doc.paragraphs:
            if KEYWORD in p.text:
                return True
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if KEYWORD in cell.text:
                        return True
    except Exception:
        pass
    return False

def search_word_docs():
    matched_by_name = []
    matched_by_content = []

    print(f"Searching for Word documents with keyword '{KEYWORD}' across drives...")

    for drive in DRIVES:
        if not os.path.exists(drive):
            continue
        print(f"Scanning drive: {drive}")
        try:
            for root, dirs, files in os.walk(drive):
                # 过滤无用大系统目录
                dirs[:] = [d for d in dirs if d not in EXCLUDE_DIRS and not d.startswith('.')]
                
                for file in files:
                    file_lower = file.lower()
                    if file_lower.endswith('.docx') or file_lower.endswith('.doc'):
                        if file.startswith('~$'): # 跳过 Office 临时锁定文件
                            continue
                            
                        filepath = os.path.join(root, file)
                        try:
                            size = os.path.getsize(filepath)
                            mtime = os.path.getmtime(filepath)
                            dt_str = datetime.datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M')

                            if KEYWORD in file:
                                matched_by_name.append({
                                    'path': filepath,
                                    'size': size,
                                    'mtime': dt_str,
                                    'match': '文件名匹配'
                                })
                            elif check_content_contains(filepath):
                                matched_by_content.append({
                                    'path': filepath,
                                    'size': size,
                                    'mtime': dt_str,
                                    'match': '正文内容匹配'
                                })
                        except Exception:
                            pass
        except Exception as e:
            print(f"Error scanning {drive}: {e}")

    total_matches = matched_by_name + matched_by_content
    print(f"\nSearch complete! Found {len(total_matches)} matching Word documents.")
    
    print("\n=== MATCHED BY FILENAME ===")
    for item in matched_by_name[:30]:
        print(f"[{item['mtime']}] {format_size(item['size'])} | {item['path']}")

    if len(matched_by_name) > 30:
        print(f"... and {len(matched_by_name) - 30} more filename matches.")

    print("\n=== MATCHED BY CONTENT ===")
    for item in matched_by_content[:30]:
        print(f"[{item['mtime']}] {format_size(item['size'])} | {item['path']}")

    result_data = {
        'total': len(total_matches),
        'by_name': matched_by_name,
        'by_content': matched_by_content
    }
    
    import json
    with open('find_docx_results.json', 'w', encoding='utf-8') as f:
        json.dump(result_data, f, ensure_ascii=False, indent=2)

if __name__ == "__main__":
    search_word_docs()
