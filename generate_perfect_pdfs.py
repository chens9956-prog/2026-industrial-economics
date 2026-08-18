import os
import json
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import win32com.client

# Exact colors from the screenshot
MAIN_BLUE_HEX = '205A84'
MAIN_BLUE = RGBColor(32, 90, 132) # #205A84
LIGHT_BLUE_HEX = 'F0F4F8'
TEXT_DARK = RGBColor(51, 51, 51)
TEXT_WHITE = RGBColor(255, 255, 255)
FOOTER_GRAY = RGBColor(153, 153, 153)

def set_cell_background(cell, fill_color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.first_child_found_in("w:shd")
    if shd is not None:
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top, bottom, start, end):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in zip(['top', 'bottom', 'left', 'right'], [top, bottom, start, end]):
        if val is not None:
            node = OxmlElement(f'w:{edge}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV', 'left', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_run_font(run, font_name, font_size, color, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in('w:tblBorders')
    if tblBorders is not None:
        tblPr.remove(tblBorders)

def extract_answers():
    transcript_path = r"C:\Users\ausu\.gemini\antigravity\brain\48a8be4f-7705-48d3-9456-30d2f3f49048\.system_generated\logs\transcript.jsonl"
    with open(transcript_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    ans_list = []
    for line in lines:
        try:
            data = json.loads(line)
            if data.get('type') == 'MCP_TOOL':
                content = data.get('content', '')
                ans = ''
                if '{"status":"success"' in content and '"answer"' in content:
                    ans = json.loads(content[content.find('{"status":"success"'):]).get('answer', '')
                elif 'The output was large and was saved to: file://' in content:
                    file_path = content.split('file:///')[1].strip().replace('/', '\\')
                    if os.path.exists(file_path):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            file_content = f.read()
                            if '{"status":"success"' in file_content and '"answer"' in file_content:
                                ans = json.loads(file_content[file_content.find('{"status":"success"'):]).get('answer', '')
                if ans:
                    ans_list.append(ans)
        except:
            pass

    answers = {}
    last_10 = ans_list[-10:]
    for i, a in enumerate(last_10):
        answers[i + 1] = a
    return answers

def create_pdfs():
    output_dir = r"I:\4产业经济学\完美复刻版_要点PDF"
    os.makedirs(output_dir, exist_ok=True)
    
    answers = extract_answers()
    if not answers:
        print("未找到任何摘要结果。")
        return
        
    print(f"找到 {len(answers)} 个章节的摘要。")
        
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    for ch_num, answer in answers.items():
        try:
            doc = Document()
            
            # Setup Page layout (A4)
            section = doc.sections[0]
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            
            ch_str = f"CH {ch_num:02d}"
            
            # 1. Header Block (Solid Dark Blue)
            head_table = doc.add_table(rows=1, cols=1)
            head_table.autofit = False
            head_table.columns[0].width = Inches(6.67)
            h_cell = head_table.cell(0, 0)
            set_cell_background(h_cell, MAIN_BLUE_HEX)
            set_cell_margins(h_cell, 200, 200, 100, 100) # add padding
            
            # Small text "产业经济学 · 第3版 · 刘志彪"
            p_small = h_cell.paragraphs[0]
            p_small.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_small.paragraph_format.space_after = Pt(2)
            r_small = p_small.add_run("产业经济学 · 第3版 · 刘志彪")
            set_run_font(r_small, 'Microsoft YaHei', 9, TEXT_WHITE, bold=False)
            
            # Large text "CH XX"
            p_large = h_cell.add_paragraph()
            p_large.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_large.paragraph_format.space_after = Pt(0)
            r_large = p_large.add_run(ch_str)
            set_run_font(r_large, 'Microsoft YaHei', 20, TEXT_WHITE, bold=True)
            
            # Add some spacing
            doc.add_paragraph().paragraph_format.space_after = Pt(12)
            
            # 2. Sub-header "知识点摘要"
            sub_table = doc.add_table(rows=1, cols=1)
            s_cell = sub_table.cell(0, 0)
            set_cell_border(s_cell, bottom={"sz": 16, "val": "single", "color": MAIN_BLUE_HEX, "space": "0"})
            set_cell_margins(s_cell, 0, 50, 0, 0)
            p_sub = s_cell.paragraphs[0]
            r_sub = p_sub.add_run("知识点摘要")
            set_run_font(r_sub, 'Microsoft YaHei', 12, MAIN_BLUE, bold=True)
            
            doc.add_paragraph().paragraph_format.space_after = Pt(8)
            
            # Extract points
            points = []
            for p_text in answer.split('\n'):
                p_text = p_text.strip()
                if p_text:
                    clean_p = p_text.replace("**", "")
                    clean_p = re.sub(r'\[\s*\d+(?:\s*,\s*\d+)*\s*\]', '', clean_p)
                    points.append(clean_p)
                    
            final_points = []
            for point in points:
                m = re.match(r"^(\d+)\.\s*(.*)", point)
                if m:
                    final_points.append((m.group(1).zfill(2), m.group(2)))
                else:
                    final_points.append(("", point))
                    
            # 3. Main Data Table
            main_table = doc.add_table(rows=len(final_points), cols=2)
            main_table.autofit = False
            main_table.columns[0].width = Inches(0.6)
            main_table.columns[1].width = Inches(6.07)
            
            remove_table_borders(main_table)
            
            for i, (num_str, content_str) in enumerate(final_points):
                row = main_table.rows[i]
                c0 = row.cells[0]
                c1 = row.cells[1]
                
                # Col 0 (Number) styling
                set_cell_background(c0, MAIN_BLUE_HEX)
                set_cell_margins(c0, 100, 100, 0, 0)
                # White thin bottom border to separate blue blocks
                set_cell_border(c0, bottom={"sz": 4, "val": "single", "color": "FFFFFF", "space": "0"})
                
                p0 = c0.paragraphs[0]
                p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p0.paragraph_format.space_after = Pt(0)
                # Vertical center is implicit or we can just use paragraph spacing
                if num_str == "":
                    num_str = f"{i+1:02d}"
                r0 = p0.add_run(num_str)
                set_run_font(r0, 'Microsoft YaHei', 11, TEXT_WHITE, bold=True)
                
                # Col 1 (Content) styling
                bg_color = 'FFFFFF' if i % 2 == 0 else LIGHT_BLUE_HEX
                set_cell_background(c1, bg_color)
                set_cell_margins(c1, 100, 100, 150, 100) # left padding 150
                set_cell_border(c1, bottom={"sz": 4, "val": "single", "color": "FFFFFF", "space": "0"})
                
                p1 = c1.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p1.paragraph_format.space_after = Pt(0)
                p1.paragraph_format.line_spacing = 1.2
                r1 = p1.add_run(content_str)
                set_run_font(r1, 'Microsoft YaHei', 10.5, TEXT_DARK, bold=False)

            # 4. Footer
            doc.add_paragraph().paragraph_format.space_after = Pt(12)
            p_foot = doc.add_paragraph()
            p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_foot = p_foot.add_run(f"CH {ch_num:02d} · 知识点摘要 · 共 {len(final_points)} 个知识点")
            set_run_font(r_foot, 'Microsoft YaHei', 9, FOOTER_GRAY, bold=False)
            
            docx_filename = f"产业经济学CH{ch_num:02d}_要点汇编_完美复刻.docx"
            pdf_filename = f"产业经济学CH{ch_num:02d}_要点汇编_完美复刻.pdf"
            
            docx_path = os.path.join(output_dir, docx_filename)
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            doc.save(docx_path)
            
            abs_docx_path = os.path.abspath(docx_path)
            abs_pdf_path = os.path.abspath(pdf_path)
            
            word_doc = word.Documents.Open(abs_docx_path)
            word_doc.SaveAs(abs_pdf_path, FileFormat=17)
            word_doc.Close()
            
            os.remove(abs_docx_path)
            
            print(f"Successfully generated {pdf_filename}")
            
        except Exception as e:
            print(f"Failed to process CH{ch_num:02d}: {str(e)}")

    word.Quit()
    print("All PDFs processed.")

if __name__ == "__main__":
    create_pdfs()
