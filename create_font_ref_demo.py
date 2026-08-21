import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_run_font(run, ascii_font="Times New Roman", east_asia_font="SimSun"):
    run.font.name = ascii_font
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), east_asia_font)
    rPr.append(rFonts)

def generate_demo_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. 标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("春")
    set_run_font(r_title, ascii_font="Times New Roman", east_asia_font="SimSun")
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(14)
    r_author = p_author.add_run("作者：朱自清 (Zhu Ziqing, 1898–1948)")
    set_run_font(r_author, ascii_font="Times New Roman", east_asia_font="SimSun")
    r_author.font.size = Pt(11)
    r_author.font.italic = True
    r_author.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 规则展示卡片
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F2F5F8")
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    rt = p.add_run("💡 全局新规生效实时测试说明：\n")
    set_run_font(rt, ascii_font="Times New Roman", east_asia_font="SimSun")
    rt.font.bold = True
    rt.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    rc = p.add_run("1. 字体配置：中文统一采用 宋体 (SimSun)，英文与数字统一采用 Times New Roman（如 2026 年、Page 12）。\n2. 标点符号：正文统一采用全角中文标点与双弯引号“与”；文末参考文献标点除外。")
    set_run_font(rc, ascii_font="Times New Roman", east_asia_font="SimSun")
    rc.font.size = Pt(9.5)
    rc.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph() # 空行

    # 正文
    p_body1 = doc.add_paragraph()
    p_body1.paragraph_format.line_spacing = 1.25
    r1 = p_body1.add_run("盼望着，盼望着，东风来了，春天的脚步近了（2026年最新修订版）。一切都像刚睡醒的样子，欣欣然张开了眼。山朗润起来了，水涨起来了，太阳的脸红起来了。")
    set_run_font(r1, ascii_font="Times New Roman", east_asia_font="SimSun")

    p_body2 = doc.add_paragraph()
    p_body2.paragraph_format.line_spacing = 1.25
    r2 = p_body2.add_run("“吹面不寒杨柳风”，不错的，像母亲的手抚摸着你。风里带来些新翻的泥土的气息，混着青草味儿，还有各种花的香，都在微微润湿的空气里孵化。桃树、杏树、梨树，你不让我，我不让你，都开满了花“赶趟儿”。")
    set_run_font(r2, ascii_font="Times New Roman", east_asia_font="SimSun")

    p_body3 = doc.add_paragraph()
    p_body3.paragraph_format.line_spacing = 1.25
    r3 = p_body3.add_run("“一年之计在于春”，刚起头儿，有的是功夫，有的是希望。春天像刚落地的娃娃，从头到脚都是新的，它生长着（Spring 2026 Report）。")
    set_run_font(r3, ascii_font="Times New Roman", east_asia_font="SimSun")

    doc.add_paragraph() # 空行

    # 文末参考文献 (References - 标点例外测试)
    h_ref = doc.add_heading(level=2)
    hr = h_ref.add_run("参考文献 (References)")
    set_run_font(hr, ascii_font="Times New Roman", east_asia_font="SimSun")
    hr.font.size = Pt(12)
    hr.font.bold = True
    hr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    refs = [
        "[1] 朱自清. 朱自清散文全集 [M]. 北京: 人民文学出版社, 1998: 45-48.",
        "[2] Smith, J. A., & Brown, L. M. (2024). Modern Chinese Literature and Epigenetic Symbolism. Journal of Asian Studies, 58(3), 112-128. https://doi.org/10.1016/j.jas.2024.05.002",
        "[3] 高建刚. 产业经济学理论与实践 [M]. 北京: 高等教育出版社, 2026: 102-115."
    ]

    for ref in refs:
        rp = doc.add_paragraph()
        rp.paragraph_format.space_before = Pt(2)
        rp.paragraph_format.space_after = Pt(4)
        rr = rp.add_run(ref)
        set_run_font(rr, ascii_font="Times New Roman", east_asia_font="SimSun")
        rr.font.size = Pt(9.5)
        rr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    output_path = "朱自清_春_字体标点与参考文献示范文档.docx"
    doc.save(output_path)
    print(f"Successfully generated font and reference demo document: {output_path}")

if __name__ == "__main__":
    generate_demo_doc()
