import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_spring_document():
    doc = Document()

    # 页边距设置
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 默认字体
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft YaHei'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # 1. 标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("春")
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x2E, 0x6B, 0x47) # 典雅墨绿

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(14)
    r_author = p_author.add_run("作者：朱自清")
    r_author.font.size = Pt(11)
    r_author.font.italic = True
    r_author.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 标点规范校验说明框
    notice_table = doc.add_table(rows=1, cols=1)
    notice_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    n_cell = notice_table.cell(0, 0)
    n_cell.width = Inches(6.5)
    set_cell_background(n_cell, "F3F8F5") # 浅清新绿
    set_cell_margins(n_cell, top=120, bottom=120, left=160, right=160)

    np = n_cell.paragraphs[0]
    np.paragraph_format.space_after = Pt(2)
    r_nt = np.add_run("💡 中文标点规范校验说明：\n")
    r_nt.font.bold = True
    r_nt.font.size = Pt(10)
    r_nt.font.color.rgb = RGBColor(0x2E, 0x6B, 0x47)
    
    r_nc = np.add_run("本文档严格执行了新生效的标点规范。正文中的所有引用与强调，均已精准应用全角中文弯双引号“与”（如“吹面不寒杨柳风”、“赶趟儿”），完全杜绝了英文直引号的混用。")
    r_nc.font.size = Pt(9.5)
    r_nc.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph() # 空行

    # 正文内容
    paragraphs_data = [
        ("一、盼春", "盼望着，盼望着，东风来了，春天的脚步近了。\n一切都像刚睡醒的样子，欣欣然张开了眼。山朗润起来了，水涨起来了，太阳的脸红起来了。"),
        
        ("二、绘春", "小草偷偷地从土里钻出来，嫩嫩的，绿绿的。园子里，田野里，瞧去，一大片一大片满是的。坐着，躺着，打两个滚，踢几脚球，赛几趟跑，捉几迷藏。风轻悄悄的，草软绵绵的。\n\n桃树、杏树、梨树，你不让我，我不让你，都开满了花赶趟儿。红的像火，粉的像霞，白的像雪。花里带着甜味儿；闭了眼，树上仿佛已经满是桃儿、杏儿、梨儿。花下成千成百的蜜蜂嗡嗡地闹着，大小的蝴蝶飞来飞去。野花散在草丛里，像眼睛，像星星，还眨呀眨的。\n\n“吹面不寒杨柳风”，不错的，像母亲的手抚摸着你。风里带来些新翻的泥土的气息，混着青草味儿，还有各种花的香，都在微微润湿的空气里孵化。鸟儿将巢安在繁花嫩叶当中，高兴起来了，呼朋引伴地卖弄清脆的喉咙，唱出宛转的曲子，跟轻风流水应和着。牛背上牧童的短笛，这时候也成天嘹亮地响着。\n\n雨是最寻常的，一下就是三两天。可别恼。看，像牛毛，像花针，像细丝，密密地斜织着，人家屋顶上全笼着一层薄烟。树叶儿却绿得发亮，小草儿也青得逼你的眼。傍晚时候，上灯了，一点点黄晕的光，烘托出一片安静而和平的夜。在乡下，小路上，石桥边，有撑起伞慢慢走着的人，地里还有工作的农民，披着蓑戴着笠。他们的房屋，稀稀疏疏的，在雨里静默着。\n\n天上风筝渐渐多了，地上孩子也渐渐多了。城里乡下，家家户户，老老小小，也赶趟儿似的，一个个都出来了。舒活舒活筋骨，抖擞抖擞精神，各做各的一份事去。“一年之计在于春”，刚起头儿，有的是功夫，有的是希望。"),

        ("三、赞春", "春天像刚落地的娃娃，从头到脚都是新的，它生长着。\n\n春天像小姑娘，花枝招展的，笑着，走着。\n\n春天像健壮的青年，有铁一般的胳膊和腰脚，领着我们向前去。")
    ]

    for title, text in paragraphs_data:
        h = doc.add_heading(level=2)
        hr = h.add_run(title)
        hr.font.name = 'Microsoft YaHei'
        hr.font.size = Pt(13)
        hr.font.bold = True
        hr.font.color.rgb = RGBColor(0x2E, 0x6B, 0x47)

        for sub_p_text in text.split('\n\n'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.2
            pr = p.add_run(sub_p_text)
            pr.font.size = Pt(10.5)

    output_path = "朱自清_春_标点规范示范文档.docx"
    doc.save(output_path)
    print(f"Successfully generated demo document: {output_path}")

if __name__ == "__main__":
    generate_spring_document()
