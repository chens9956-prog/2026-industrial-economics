import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

doc = Document()

# 设置页边距
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# 设置默认字体
normal_style = doc.styles['Normal']
normal_style.font.name = 'Microsoft YaHei'
normal_style.font.size = Pt(10.5)
normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# 标题
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("NotebookLM 高效研究提示词指令集\n（中英文对照版）")
title_run.font.name = 'Microsoft YaHei'
title_run.font.size = Pt(20)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # 优雅深蓝

# 前言说明
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
sub_run = sub_p.add_run("本文档整理了在 Google NotebookLM 中用于学术文献分析、矛盾检索、综合矩阵构建及综述提纲生成的 6 套高阶 Prompt 指令。每条指令均提供精细的中文简体翻译与英文原文，方便即拿即用。")
sub_run.font.size = Pt(10)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph() # 空行

prompts_data = [
    {
        "id": 1,
        "title": "指令一：文献总体概览与相关性筛选",
        "category": "文献初步筛选 (Initial Screening)",
        "zh": "请用一段话简要概括这些资料来源共同涵盖的内容，并指出其中似乎与上述主题关系不大的任何资料，以便我决定是否将其从本笔记本中移除。",
        "en": "Give me a one-paragraph overview of what these sources collectively cover, and flag any of them that seem tangential to the topic above so I can decide whether to remove them from this notebook.",
        "usage": "适用于在 NotebookLM 中刚导入多篇文献时，快速评估所有文献的整体主题相关度，过滤不相关论文。"
    },
    {
        "id": 2,
        "title": "指令二：文献矛盾与观点冲突检索",
        "category": "文献矛盾对比 (Contradiction Search)",
        "zh": "请找出这些资料来源之间的任何直接矛盾之处。对于每处矛盾，请列出两篇相关论文，分别用一句话说明它们相互对立的观点，并引用支持各自观点的具体原文段落。仅报告能够附带直接原文引用的矛盾。",
        "en": "Identify any direct contradictions between these sources. For each contradiction, name both papers, state their opposing claims in one sentence each, and cite the specific passage supporting each claim. Only report contradictions you can support with a direct citation.",
        "usage": "适用于寻找不同研究之间的争议点与冲突证据，确保所有发现都有直接原文出处支撑（防止 AI 幻觉）。"
    },
    {
        "id": 3,
        "title": "指令三：文献综合矩阵表生成 (Synthesis Matrix)",
        "category": "文献综述矩阵 (Synthesis Table)",
        "zh": "请为这些论文绘制一份排版整洁、格式规范的文献综合矩阵表（Synthesis Matrix），表格列应包含：作者（Authors）、研究方法（Methods）、数据集（Datasets）、主要发现（Main Findings）、局限性（Limitations）、未来方向（Future directions）、研究空白（Research Gap）、样本量（Sample Size）以及主题（Theme）。",
        "en": "Create a literature Synthesis Matrix for these papers in a clean well formatted table with columns for :Authors, Methods, Datasets, Main Findings, Limitations, Future directions, Research Gap, Sample Size, and Theme",
        "usage": "适用于快速生成多维度的文献对比表格，方便直接导出或整理到论文的 Literature Review 章节中。"
    },
    {
        "id": 4,
        "title": "指令四：研究空白矩阵与创新点挖掘",
        "category": "研究空白与创新 (Research Gap Analysis)",
        "zh": "请充当一名研究助理。仅使用本笔记本中的资料来源，分析相关文献并梳理出：已知事实、文献间的冲突证据、方法学上的缺陷、未被探索的人群或情境，以及尚未解决的研究空白。请将上述内容总结为一份结构化的“研究空白矩阵表”（Research Gap Matrix），并在每个论点后注明具体的参考来源；最后用一段话总结我的研究可能做出的创新贡献。所有关于“研究空白”的论断必须严格基于已上传资料的覆盖盲区——切勿超出现有资料进行主观推测。",
        "en": "Act as a research assistant. Using only the sources in this notebook, analyze the literature and identify: what is already known, conflicting evidence between sources, methodological weaknesses, unexplored populations or contexts, and the remaining research gap Summarize this as a structured Research Gap Matrix citing the specific source for every claim, and conclude with a one-paragraph statement of the novel contribution my research could make. Base every gap claim strictly on the absence of coverage across these specific sources -- do not speculate beyond what's uploaded.",
        "usage": "适用于撰写论文开题报告、申请基金或论证研究创新性 (Novelty) 时，系统梳理当前研究不足并提出独到贡献。"
    },
    {
        "id": 5,
        "title": "指令五：主题文献综述高度详细提纲生成",
        "category": "综述提纲生成 (Literature Review Outline)",
        "zh": "请为题为《人类疾病的表观遗传学图谱：致病机制、生物标志物与下一代疗法》的文献综述生成一份结构化且高度详细的提纲。仅使用所提供的资料来源，构建提纲的前半部分以界定核心表观遗传机制及其在癌症、自身免疫性疾病、神经退行性疾病和心血管疾病中的致病作用；提纲的后半部分则用于评估新兴的临床应用，重点关注诊断性生物标志物以及下一代疗法（如基于 CRISPR 的精准表观基因组编辑和药理学抑制剂）。",
        "en": "Generate a structured, highly detailed outline for a literature review titled 'The Epigenetic Landscape of Human Disease: Pathogenic Mechanisms, Biomarkers, and Next-Generation Therapeutics: Using only the provided sources, organize the outline to first define core epigenetic mechanisms and their pathogenic roles across cancer, autoimmune, neurodegenerative, and cardiovascular diseases. Finally, structure the latter half of the outline to evaluate emerging clinical applications, specifically focusing on diagnostic biomarkers and next-generation therapeutics like CRISPR-based precision epigenome editing and pharmacological inhibitors.\"",
        "usage": "适用于搭建高水平综述论文（Review Paper）的大纲与框架（可根据自己的具体论文题目修改标题与疾病类型）。"
    },
    {
        "id": 6,
        "title": "指令六：引言部分的编年史叙述撰写",
        "category": "编年学术演进叙述 (Chronological Narrative)",
        "zh": "请充当一名研究助理，帮我撰写一段关于人类基因表观遗传调控研究演进历程的按时间顺序叙述的文字，用于论文的前言/引言部分。仅使用现有的资料来源，按照出版年份顺序整理关键进展、重要发现以及学术观点的演变。对于每个时期或里程碑节点，请说明具体年份或年份范围、推动该进展的论文、在此之前的学界认知，以及随之发生的变化。请以流畅通顺的段落式叙述呈现（而非表格形式），以便直接修改并融入引言部分。每个论断均须包含文内引用。",
        "en": "Act as a research assistant helping me draft a chronological narrative of how research on epigenetic regulation of human genes has developed over time, for use in my introduction section. Using only these sources, rganize the key developments, findings, and shifts in thinking in chronological order by publication year. For each period or milestone, state the year or range, which paper(s) drove that development, what the field understood before that point, and what changed as a result. Write this as a flowing narrative paragraph, not a table, suitable for adapting into an introduction section. Include an inline citation for every claim.",
        "usage": "适用于撰写论文 Background / Introduction 中关于“研究历史演进与前人发展脉络”的流畅叙事段落。"
    }
]

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

for item in prompts_data:
    # 标题 (Heading 2)
    h = doc.add_heading(level=2)
    h_run = h.add_run(item["title"])
    h_run.font.name = 'Microsoft YaHei'
    h_run.font.size = Pt(13)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    # 建立中英文对照表格 (1 列，单单元格或双单元格)
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # 第一行：中文简体翻译
    cell_zh = table.cell(0, 0)
    cell_zh.width = Inches(6.5)
    set_cell_background(cell_zh, "F2F5F8") # 优雅浅蓝灰
    set_cell_margins(cell_zh, top=120, bottom=120, left=180, right=180)
    
    p_zh = cell_zh.paragraphs[0]
    p_zh.paragraph_format.space_after = Pt(2)
    r_zh_label = p_zh.add_run("💡 中文简体翻译 (Prompt)：\n")
    r_zh_label.font.bold = True
    r_zh_label.font.size = Pt(10)
    r_zh_label.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    r_zh = p_zh.add_run(item["zh"])
    r_zh.font.size = Pt(10.5)
    r_zh.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    
    # 第二行：英文原文
    cell_en = table.cell(1, 0)
    cell_en.width = Inches(6.5)
    set_cell_background(cell_en, "FAFAFA") # 浅灰
    set_cell_margins(cell_en, top=120, bottom=120, left=180, right=180)
    
    p_en = cell_en.paragraphs[0]
    p_en.paragraph_format.space_after = Pt(2)
    r_en_label = p_en.add_run("🔤 英文原文 (Original Prompt)：\n")
    r_en_label.font.bold = True
    r_en_label.font.size = Pt(10)
    r_en_label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    r_en = p_en.add_run(item["en"])
    r_en.font.size = Pt(10)
    r_en.font.name = 'Consolas'
    r_en.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # 使用场景说明
    p_usage = doc.add_paragraph()
    p_usage.paragraph_format.space_before = Pt(4)
    p_usage.paragraph_format.space_after = Pt(14)
    r_u_title = p_usage.add_run("📌 应用场景建议：")
    r_u_title.font.bold = True
    r_u_title.font.size = Pt(9.5)
    r_u_title.font.color.rgb = RGBColor(0x8C, 0x6B, 0x00)
    
    r_u = p_usage.add_run(item["usage"])
    r_u.font.size = Pt(9.5)
    r_u.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

output_path = "NotebookLM_提示词指令集_中英文对照.docx"
doc.save(output_path)
print(f"Successfully generated {output_path}")
